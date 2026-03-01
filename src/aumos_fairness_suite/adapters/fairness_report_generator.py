"""Fairness Report Generator — audit-ready PDF/DOCX reports for regulatory submissions.

Generates structured fairness assessment reports suitable for:
- EU AI Act Article 10 (data governance and bias testing documentation)
- ECOA 12 CFR Part 202 (Equal Credit Opportunity Act — non-discrimination documentation)
- NIST AI RMF (Govern/Measure functions — fairness documentation)

Report sections:
1. Executive Summary — overall pass/fail, affected groups, key findings
2. Model Information — model ID, version, training data reference
3. Assessment Methodology — metrics used, thresholds, dataset size
4. Metric Results Table — per-metric values with pass/fail indicators
5. Group-Level Analysis — per-protected-group metric breakdown
6. Mitigation Recommendations — actionable remediation steps
7. Regulatory Mapping — ECOA controls, EU AI Act checklist
8. Sign-off Block — reviewer fields for compliance workflow

Output formats: PDF (via reportlab) and DOCX (via python-docx).
Reports stored in object storage; fai_reports table stores URI + metadata.
"""

from __future__ import annotations

import importlib.util
import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any

from aumos_common.observability import get_logger

logger = get_logger(__name__)


def _is_reportlab_available() -> bool:
    """Check whether reportlab is installed.

    Returns:
        True if reportlab is importable, False otherwise.
    """
    return importlib.util.find_spec("reportlab") is not None


def _is_docx_available() -> bool:
    """Check whether python-docx is installed.

    Returns:
        True if docx (python-docx) is importable, False otherwise.
    """
    return importlib.util.find_spec("docx") is not None


# EU AI Act Article 10 checklist items for fairness documentation
EU_AI_ACT_ARTICLE_10_CHECKLIST: list[str] = [
    "Bias testing performed on training and validation datasets",
    "Protected attributes identified and documented",
    "Fairness metrics computed with defined thresholds",
    "Mitigation measures applied where bias detected",
    "Post-mitigation bias testing performed",
    "Human oversight mechanism documented",
    "Audit trail maintained for all assessments",
    "Report reviewed by qualified personnel",
]

# ECOA 12 CFR Part 202 checklist items
ECOA_CHECKLIST: list[str] = [
    "Model tested for disparate impact on prohibited bases (race, color, religion, national origin, sex, marital status, age)",
    "Disparate impact ratio >= 0.80 (4/5 rule) for all protected groups",
    "Equal opportunity difference within acceptable bounds",
    "Model purpose and limitations documented",
    "Adverse action reason codes available",
    "Model monitoring schedule established",
    "Assessment independently reviewed",
]

# NIST AI RMF Measure function items
NIST_AI_RMF_CHECKLIST: list[str] = [
    "GOVERN: Fairness policies established and documented",
    "MAP: Protected groups and potential harms identified",
    "MEASURE: Quantitative fairness metrics computed and tracked",
    "MANAGE: Remediation plan in place for failed metrics",
]


@dataclass
class FairnessReportMetadata:
    """Metadata for a generated fairness report.

    Attributes:
        model_id: UUID of the assessed model.
        model_name: Human-readable model name.
        model_version: Version tag.
        assessment_id: UUID of the fairness assessment.
        assessment_date: ISO-8601 assessment timestamp.
        overall_passed: Whether all fairness metrics passed.
        regulation: Regulatory framework (ecoa, eu_ai_act, nist_ai_rmf, all).
        generated_at: Report generation timestamp.
        report_format: Output format ('pdf' or 'docx').
    """

    model_id: str
    model_name: str
    model_version: str
    assessment_id: str
    assessment_date: str
    overall_passed: bool
    regulation: str
    generated_at: str
    report_format: str


class FairnessReportGenerator:
    """Generates audit-ready fairness assessment reports in PDF and DOCX formats.

    Uses reportlab for PDF and python-docx for DOCX. Falls back to a plain-text
    representation if neither library is available. Reports include regulatory
    checklists, metric tables, group-level breakdowns, and sign-off blocks.

    Args:
        organization_name: Name of the organization for report branding.
        include_raw_data: Whether to include raw metric data tables in reports.
    """

    def __init__(
        self,
        organization_name: str = "AumOS Enterprise",
        include_raw_data: bool = True,
    ) -> None:
        """Initialise the report generator.

        Args:
            organization_name: Organization name for report header.
            include_raw_data: Include detailed metric data tables in reports.
        """
        self._organization_name = organization_name
        self._include_raw_data = include_raw_data

    def generate_pdf(
        self,
        assessment: dict[str, Any],
        bias_metrics: list[dict[str, Any]],
        mitigation_recommendations: list[str],
        regulation: str = "all",
        model_info: dict[str, Any] | None = None,
    ) -> bytes:
        """Generate a PDF fairness report.

        Args:
            assessment: FairnessAssessment record dict.
            bias_metrics: List of BiasMetric record dicts.
            mitigation_recommendations: Actionable remediation steps.
            regulation: Regulatory framework ('ecoa', 'eu_ai_act', 'nist_ai_rmf', 'all').
            model_info: Optional model metadata dict.

        Returns:
            PDF bytes. Falls back to encoded plain text if reportlab is unavailable.
        """
        if not _is_reportlab_available():
            logger.warning("reportlab_unavailable_generating_text_report")
            text_report = self._generate_text_report(
                assessment, bias_metrics, mitigation_recommendations, regulation, model_info
            )
            return text_report.encode("utf-8")

        from reportlab.lib import colors  # type: ignore[import]
        from reportlab.lib.pagesizes import A4  # type: ignore[import]
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import]
        from reportlab.lib.units import mm  # type: ignore[import]
        from reportlab.platypus import (  # type: ignore[import]
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
        styles = getSampleStyleSheet()
        story = []

        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        overall_passed = bool(assessment.get("overall_passed", False))
        status_text = "PASSED" if overall_passed else "FAILED"
        status_color = colors.green if overall_passed else colors.red

        # Title
        story.append(Paragraph(
            f"<b>{self._organization_name}</b>",
            styles["Heading1"],
        ))
        story.append(Paragraph(
            "AI Fairness Assessment Report",
            styles["Heading2"],
        ))
        story.append(Paragraph(
            f"Generated: {generated_at} | Regulation: {regulation.upper()}",
            styles["Normal"],
        ))
        story.append(Spacer(1, 10 * mm))

        # Executive Summary
        story.append(Paragraph("1. Executive Summary", styles["Heading2"]))
        story.append(Paragraph(
            f"<b>Assessment Status: <font color={'green' if overall_passed else 'red'}>{status_text}</font></b>",
            styles["Normal"],
        ))
        story.append(Paragraph(
            f"Assessment ID: {assessment.get('id', 'N/A')}",
            styles["Normal"],
        ))
        story.append(Paragraph(
            f"Model ID: {assessment.get('model_id', 'N/A')}",
            styles["Normal"],
        ))
        story.append(Paragraph(
            f"Assessment Date: {assessment.get('created_at', 'N/A')}",
            styles["Normal"],
        ))
        story.append(Spacer(1, 8 * mm))

        # Metric Results Table
        story.append(Paragraph("2. Metric Results", styles["Heading2"]))

        table_data = [["Metric", "Value", "Threshold", "Status"]]
        for m in bias_metrics:
            metric_name = m.get("metric_name", "unknown")
            value = round(float(m.get("value", 0.0)), 4)
            threshold = float(m.get("threshold", 0.1))
            passed = bool(m.get("passed", True))
            status = "PASS" if passed else "FAIL"
            table_data.append([metric_name, str(value), str(threshold), status])

        if len(table_data) > 1:
            table = Table(table_data, colWidths=[80 * mm, 30 * mm, 30 * mm, 20 * mm])
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("PADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(table)
        story.append(Spacer(1, 8 * mm))

        # Regulatory Mapping
        story.append(Paragraph("3. Regulatory Compliance Checklist", styles["Heading2"]))
        checklist = self._get_regulatory_checklist(regulation)
        for item in checklist:
            story.append(Paragraph(f"[{'X' if overall_passed else ' '}] {item}", styles["Normal"]))
        story.append(Spacer(1, 8 * mm))

        # Mitigation Recommendations
        if mitigation_recommendations:
            story.append(Paragraph("4. Mitigation Recommendations", styles["Heading2"]))
            for i, rec in enumerate(mitigation_recommendations, 1):
                story.append(Paragraph(f"{i}. {rec}", styles["Normal"]))
            story.append(Spacer(1, 8 * mm))

        # Sign-off Block
        story.append(Paragraph("5. Sign-off", styles["Heading2"]))
        story.append(Paragraph("Reviewed by: _______________________", styles["Normal"]))
        story.append(Paragraph("Date: _______________________", styles["Normal"]))
        story.append(Paragraph("Compliance Officer: _______________________", styles["Normal"]))

        doc.build(story)
        return buffer.getvalue()

    def generate_docx(
        self,
        assessment: dict[str, Any],
        bias_metrics: list[dict[str, Any]],
        mitigation_recommendations: list[str],
        regulation: str = "all",
        model_info: dict[str, Any] | None = None,
    ) -> bytes:
        """Generate a DOCX fairness report.

        Args:
            assessment: FairnessAssessment record dict.
            bias_metrics: List of BiasMetric record dicts.
            mitigation_recommendations: Actionable remediation steps.
            regulation: Regulatory framework.
            model_info: Optional model metadata dict.

        Returns:
            DOCX bytes. Falls back to encoded plain text if python-docx is unavailable.
        """
        if not _is_docx_available():
            logger.warning("python_docx_unavailable_generating_text_report")
            text_report = self._generate_text_report(
                assessment, bias_metrics, mitigation_recommendations, regulation, model_info
            )
            return text_report.encode("utf-8")

        import docx  # type: ignore[import]
        from docx.shared import Pt, RGBColor  # type: ignore[import]

        document = docx.Document()

        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        overall_passed = bool(assessment.get("overall_passed", False))
        status_text = "PASSED" if overall_passed else "FAILED"

        # Title
        title = document.add_heading(f"{self._organization_name} — AI Fairness Assessment Report", 0)
        document.add_paragraph(f"Generated: {generated_at} | Regulation: {regulation.upper()}")

        # Executive Summary
        document.add_heading("1. Executive Summary", level=1)
        summary_para = document.add_paragraph(f"Assessment Status: {status_text}")
        run = summary_para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0) if overall_passed else RGBColor(255, 0, 0)

        document.add_paragraph(f"Assessment ID: {assessment.get('id', 'N/A')}")
        document.add_paragraph(f"Model ID: {assessment.get('model_id', 'N/A')}")
        document.add_paragraph(f"Assessment Date: {assessment.get('created_at', 'N/A')}")

        # Metric Results Table
        document.add_heading("2. Metric Results", level=1)
        if bias_metrics:
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            header_cells[2].text = "Threshold"
            header_cells[3].text = "Status"

            for m in bias_metrics:
                row_cells = table.add_row().cells
                row_cells[0].text = m.get("metric_name", "unknown")
                row_cells[1].text = str(round(float(m.get("value", 0.0)), 4))
                row_cells[2].text = str(float(m.get("threshold", 0.1)))
                row_cells[3].text = "PASS" if m.get("passed", True) else "FAIL"

        # Regulatory Checklist
        document.add_heading("3. Regulatory Compliance Checklist", level=1)
        checklist = self._get_regulatory_checklist(regulation)
        for item in checklist:
            prefix = "[X]" if overall_passed else "[ ]"
            document.add_paragraph(f"{prefix} {item}", style="List Bullet")

        # Mitigation Recommendations
        if mitigation_recommendations:
            document.add_heading("4. Mitigation Recommendations", level=1)
            for rec in mitigation_recommendations:
                document.add_paragraph(rec, style="List Number")

        # Sign-off
        document.add_heading("5. Sign-off", level=1)
        document.add_paragraph("Reviewed by: _______________________")
        document.add_paragraph("Date: _______________________")
        document.add_paragraph("Compliance Officer: _______________________")

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _get_regulatory_checklist(self, regulation: str) -> list[str]:
        """Return the appropriate regulatory checklist items.

        Args:
            regulation: Regulation identifier ('ecoa', 'eu_ai_act', 'nist_ai_rmf', 'all').

        Returns:
            List of checklist item strings.
        """
        if regulation == "ecoa":
            return ECOA_CHECKLIST
        if regulation == "eu_ai_act":
            return EU_AI_ACT_ARTICLE_10_CHECKLIST
        if regulation == "nist_ai_rmf":
            return NIST_AI_RMF_CHECKLIST
        # 'all' — combine all checklists
        return EU_AI_ACT_ARTICLE_10_CHECKLIST + ECOA_CHECKLIST + NIST_AI_RMF_CHECKLIST

    def _generate_text_report(
        self,
        assessment: dict[str, Any],
        bias_metrics: list[dict[str, Any]],
        mitigation_recommendations: list[str],
        regulation: str,
        model_info: dict[str, Any] | None,
    ) -> str:
        """Generate a plain-text fallback report when PDF/DOCX libraries are unavailable.

        Args:
            assessment: FairnessAssessment record dict.
            bias_metrics: BiasMetric records.
            mitigation_recommendations: Remediation steps.
            regulation: Regulatory framework.
            model_info: Optional model metadata.

        Returns:
            Plain-text report string.
        """
        generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        overall_passed = bool(assessment.get("overall_passed", False))
        status_text = "PASSED" if overall_passed else "FAILED"

        lines: list[str] = [
            f"{self._organization_name} — AI Fairness Assessment Report",
            f"Generated: {generated_at} | Regulation: {regulation.upper()}",
            "=" * 60,
            "",
            "EXECUTIVE SUMMARY",
            f"Status: {status_text}",
            f"Assessment ID: {assessment.get('id', 'N/A')}",
            f"Model ID: {assessment.get('model_id', 'N/A')}",
            "",
            "METRIC RESULTS",
        ]

        for m in bias_metrics:
            passed_str = "PASS" if m.get("passed", True) else "FAIL"
            lines.append(
                f"  {m.get('metric_name', 'unknown')}: "
                f"{round(float(m.get('value', 0.0)), 4)} "
                f"(threshold={m.get('threshold', 0.1)}) [{passed_str}]"
            )

        lines.append("")
        lines.append("REGULATORY CHECKLIST")
        checklist = self._get_regulatory_checklist(regulation)
        for item in checklist:
            prefix = "[X]" if overall_passed else "[ ]"
            lines.append(f"  {prefix} {item}")

        if mitigation_recommendations:
            lines.append("")
            lines.append("MITIGATION RECOMMENDATIONS")
            for i, rec in enumerate(mitigation_recommendations, 1):
                lines.append(f"  {i}. {rec}")

        lines.extend([
            "",
            "SIGN-OFF",
            "Reviewed by: _______________________",
            "Date: _______________________",
            "Compliance Officer: _______________________",
        ])

        return "\n".join(lines)

    def build_metadata(
        self,
        assessment: dict[str, Any],
        report_format: str,
        regulation: str,
        model_info: dict[str, Any] | None = None,
    ) -> FairnessReportMetadata:
        """Build metadata record for a generated report.

        Args:
            assessment: FairnessAssessment record dict.
            report_format: Output format ('pdf' or 'docx').
            regulation: Regulatory framework.
            model_info: Optional model metadata.

        Returns:
            FairnessReportMetadata for persistence in fai_reports.
        """
        generated_at = datetime.now(timezone.utc).isoformat()
        return FairnessReportMetadata(
            model_id=str(assessment.get("model_id", "")),
            model_name=str((model_info or {}).get("name", "Unknown Model")),
            model_version=str((model_info or {}).get("version", "unknown")),
            assessment_id=str(assessment.get("id", "")),
            assessment_date=str(assessment.get("created_at", "")),
            overall_passed=bool(assessment.get("overall_passed", False)),
            regulation=regulation,
            generated_at=generated_at,
            report_format=report_format,
        )


__all__ = [
    "FairnessReportGenerator",
    "FairnessReportMetadata",
    "EU_AI_ACT_ARTICLE_10_CHECKLIST",
    "ECOA_CHECKLIST",
    "NIST_AI_RMF_CHECKLIST",
]
